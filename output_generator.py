"""
output_generator.py — Converts agent outputs into Word doc (CV) and PDF (full report).

CV     → python-docx (pure Python, no Node.js required)
Report → reportlab (Python)
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    PageBreak, Table, TableStyle,
)
from reportlab.lib.enums import TA_CENTER


# ── Colour palette ────────────────────────────────────────────────────────────
NAVY = colors.HexColor("#1B4F72")
BLUE = colors.HexColor("#2E86C1")
GREY = colors.HexColor("#5D6D7E")
LIGHT = colors.HexColor("#EBF5FB")
WHITE = colors.white
BLACK = colors.black

_NAVY_RGB = RGBColor(0x1B, 0x4F, 0x72)
_BLUE_RGB = RGBColor(0x2E, 0x86, 0xC1)
_GREY_RGB = RGBColor(0x5D, 0x6D, 0x7E)


# ── Word Document (CV) ────────────────────────────────────────────────────────

def _add_hr(doc: Document, color_hex: str = "2E86C1") -> None:
    """Insert a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _is_section_header(line: str) -> bool:
    """True if line looks like a CV section header (ALL CAPS or ## prefix)."""
    stripped = line.strip()
    if stripped.startswith("## "):
        return True
    # ALL CAPS word(s), at least 3 chars, no digits — e.g. EXPERIENCE, SKILLS
    return bool(re.match(r"^[A-Z][A-Z\s/]{2,}$", stripped))


def _is_contact_line(line: str) -> bool:
    """True for lines containing email/phone/linkedin patterns."""
    return bool(re.search(r"[@|•·]", line) or re.search(r"\d{5,}", line))


def _set_run_font(run, size_pt: float, bold: bool = False,
                  color: RGBColor | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def generate_cv_docx(cv_text: str, output_path: str) -> str:
    """
    Build a formatted .docx CV from plain CV text using python-docx.
    No Node.js or temp files required.
    Returns the path to the generated .docx file.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    lines = [l.rstrip() for l in cv_text.splitlines()]
    i = 0
    name_written = False

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Skip completely blank lines between sections
        if not stripped:
            i += 1
            continue

        # ── Candidate name (first non-blank line) ─────────────────────────
        if not name_written:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped)
            _set_run_font(run, 20, bold=True, color=_NAVY_RGB)
            name_written = True
            i += 1
            continue

        # ── Contact line (email / phone / LinkedIn) ────────────────────────
        if _is_contact_line(stripped) and i < 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped)
            _set_run_font(run, 9, color=_GREY_RGB)
            _add_hr(doc, "2E86C1")
            i += 1
            continue

        # ── Section header ─────────────────────────────────────────────────
        if _is_section_header(stripped):
            header_text = stripped.lstrip("# ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(header_text.upper())
            _set_run_font(run, 11, bold=True, color=_NAVY_RGB)
            _add_hr(doc, "1B4F72")
            i += 1
            continue

        # ── Company | Role | Date line ─────────────────────────────────────
        if "|" in stripped and not stripped.startswith("-"):
            parts = [p.strip() for p in stripped.split("|")]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            # Company + Role in bold navy, date right-aligned via tab
            run = p.add_run("  ".join(parts[:-1]))
            _set_run_font(run, 10, bold=True, color=_NAVY_RGB)
            if len(parts) > 1:
                run2 = p.add_run(f"  {parts[-1]}")
                _set_run_font(run2, 9, color=_GREY_RGB)
            i += 1
            continue

        # ── Bullet point ───────────────────────────────────────────────────
        if stripped.startswith("- ") or stripped.startswith("• "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(bullet_text)
            _set_run_font(run, 9.5)
            i += 1
            continue

        # ── Numbered item ──────────────────────────────────────────────────
        if re.match(r"^\d+\.", stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(stripped)
            _set_run_font(run, 9.5)
            i += 1
            continue

        # ── Plain body line ────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(stripped)
        _set_run_font(run, 9.5)
        i += 1

    doc.save(output_path)
    print(f"  CV saved → {output_path}")
    return output_path


# ── PDF Report ────────────────────────────────────────────────────────────────

def _styles() -> dict:
    base = getSampleStyleSheet()

    title = ParagraphStyle(
        "ReportTitle",
        parent=base["Title"],
        fontSize=24,
        textColor=NAVY,
        spaceAfter=6,
        fontName="Helvetica-Bold",
        alignment=TA_CENTER,
    )
    subtitle = ParagraphStyle(
        "ReportSubtitle",
        parent=base["Normal"],
        fontSize=11,
        textColor=GREY,
        spaceAfter=20,
        fontName="Helvetica",
        alignment=TA_CENTER,
    )
    h1 = ParagraphStyle(
        "SectionH1",
        parent=base["Heading1"],
        fontSize=14,
        textColor=WHITE,
        backColor=NAVY,
        spaceBefore=18,
        spaceAfter=4,
        leftIndent=-6,
        rightIndent=-6,
        fontName="Helvetica-Bold",
        borderPadding=(4, 6, 4, 6),
    )
    h2 = ParagraphStyle(
        "SectionH2",
        parent=base["Heading2"],
        fontSize=12,
        textColor=NAVY,
        spaceBefore=12,
        spaceAfter=4,
        fontName="Helvetica-Bold",
    )
    body = ParagraphStyle(
        "Body",
        parent=base["Normal"],
        fontSize=10,
        textColor=BLACK,
        spaceBefore=2,
        spaceAfter=4,
        leading=14,
        fontName="Helvetica",
    )
    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=16,
        bulletIndent=6,
        spaceBefore=1,
        spaceAfter=2,
    )
    label = ParagraphStyle(
        "Label",
        parent=body,
        textColor=GREY,
        fontSize=9,
        fontName="Helvetica-Oblique",
    )
    return {
        "title": title, "subtitle": subtitle,
        "h1": h1, "h2": h2,
        "body": body, "bullet": bullet, "label": label,
    }


def _sanitise_for_reportlab(text: str) -> str:
    """
    Make a string safe to pass to ReportLab's XML-based Paragraph renderer.
    Handles: XML special chars, markdown bold/italic, curly quotes.
    """
    # Curly quotes → straight (ReportLab's XML parser chokes on them)
    text = text.replace("‘", "'").replace("’", "'")
    text = text.replace("“", '"').replace("”", '"')
    # Em/en dashes → plain hyphen-minus
    text = text.replace("—", "--").replace("–", "-")
    # Strip markdown bold (**text** or __text__) and italic (*text* or _text_)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    # XML special characters — must be last
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return text


def _parse_markdown_to_flowables(text: str, s: dict) -> list:
    """
    Convert agent markdown output to ReportLab flowables.
    Handles ##/### headers, bullet lists, numbered lists, and plain paragraphs.
    All text is sanitised to prevent ReportLab XML parse errors.
    """
    flowables = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            flowables.append(Spacer(1, 4))
        elif stripped.startswith("### "):
            safe = _sanitise_for_reportlab(stripped[4:])
            flowables.append(Paragraph(safe, s["h2"]))
        elif stripped.startswith("## "):
            safe = _sanitise_for_reportlab(stripped[3:])
            flowables.append(Spacer(1, 6))
            flowables.append(Paragraph(safe, s["h1"]))
            flowables.append(HRFlowable(color=BLUE, thickness=1, spaceAfter=4))
        elif stripped.startswith("# "):
            safe = _sanitise_for_reportlab(stripped[2:])
            flowables.append(Paragraph(safe, s["h1"]))
        elif stripped.startswith("- ") or stripped.startswith("• "):
            safe = _sanitise_for_reportlab(stripped[2:])
            flowables.append(Paragraph(f"&bull; {safe}", s["bullet"]))
        elif re.match(r"^\d+\.\s", stripped):
            safe = _sanitise_for_reportlab(stripped)
            flowables.append(Paragraph(safe, s["bullet"]))
        else:
            safe = _sanitise_for_reportlab(stripped)
            flowables.append(Paragraph(safe, s["body"]))
    return flowables


def generate_report_pdf(
    company_name: str,
    jd_analysis: str,
    company_research: str,
    output_path: str,
) -> str:
    """
    Generate a comprehensive PDF report combining all agent outputs.
    Returns the path to the generated PDF.
    """
    s = _styles()
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    story = []
    generated = datetime.now().strftime("%B %d, %Y at %I:%M %p")

    # ── Cover / Title ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph("Job Intelligence Report", s["title"]))
    story.append(Paragraph(f"Target Company: {_sanitise_for_reportlab(company_name)}", s["subtitle"]))
    story.append(Paragraph(f"Generated: {generated}", s["label"]))
    story.append(HRFlowable(color=NAVY, thickness=2, spaceAfter=20))
    story.append(Spacer(1, 0.3 * inch))

    # ── Table of Contents ─────────────────────────────────────────────────────
    toc_data = [
        ["Section", "Content"],
        ["1", "JD Analysis — Fit Score, Keywords, Gaps"],
        ["2", "Company Intelligence — Culture, Salary, Process"],
    ]
    toc_table = Table(toc_data, colWidths=[0.5 * inch, 6 * inch])
    toc_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D5D8DC")),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(toc_table)
    story.append(PageBreak())

    # ── Section 1: JD Analysis ────────────────────────────────────────────────
    story.append(Paragraph("SECTION 1 — JD ANALYSIS", s["h1"]))
    story.append(HRFlowable(color=BLUE, thickness=1, spaceAfter=8))
    story += _parse_markdown_to_flowables(jd_analysis, s)
    story.append(PageBreak())

    # ── Section 2: Company Intelligence ──────────────────────────────────────
    story.append(Paragraph("SECTION 2 — COMPANY INTELLIGENCE", s["h1"]))
    story.append(HRFlowable(color=BLUE, thickness=1, spaceAfter=8))
    story += _parse_markdown_to_flowables(company_research, s)
    story.append(PageBreak())

    doc.build(story)
    print(f"  Report saved → {output_path}")
    return output_path
